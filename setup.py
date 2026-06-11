#!/usr/bin/env python3
"""
setup.py — One-time developer onboarding for CV Builder

Reads a developer's existing resume (PDF or DOCX), extracts their profile
using Claude API, and generates a personal config folder containing:
  _profile.md
  _resume_format.md
  _style_a.css
  _style_coverletter.css
  GenerateResumePrompt.md

Usage:
  python setup.py <resume.pdf|resume.docx> --output <config_folder>
  python setup.py JohnDoe_Resume.pdf --output "C:/Users/John/cv_config"

Requirements:
  pip install anthropic pdfplumber python-docx
  ANTHROPIC_API_KEY environment variable must be set.
"""

import sys
import os
import re
import shutil
import argparse
from pathlib import Path
from dotenv import load_dotenv

load_dotenv(Path(__file__).parent / '.env')

try:
    import pdfplumber
except ImportError:
    print("ERROR: pdfplumber not installed. Run: pip install pdfplumber")
    sys.exit(1)

try:
    import docx
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

try:
    import anthropic
except ImportError:
    print("ERROR: anthropic not installed. Run: pip install anthropic")
    sys.exit(1)


# ─── TEXT EXTRACTION ─────────────────────────────────────────────────────────

def extract_pdf(path: Path) -> str:
    text = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text.append(t)
    return '\n\n'.join(text)


def extract_docx(path: Path) -> str:
    doc  = docx.Document(str(path))
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append('  |  '.join(cells))
    return '\n'.join(lines)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == '.pdf':
        print('  Extracting text from PDF...')
        return extract_pdf(path)
    elif suffix in ('.docx', '.doc'):
        print('  Extracting text from DOCX...')
        return extract_docx(path)
    else:
        print(f'ERROR: Unsupported file type: {suffix}. Use PDF or DOCX.')
        sys.exit(1)


# ─── CLAUDE API ──────────────────────────────────────────────────────────────

PROFILE_PROMPT = """You are a professional resume parser.

Given the raw text extracted from a developer's resume below, generate a structured _profile.md file.

Follow this EXACT format:

# {Full Name} — Master Profile

## Personal Info (fixed — never change)
- **Full name:** {Full Name}
- **Total experience: X+ years** ({start year} – present — never reduce or change this number)
- **Location:** {City, State/Country}
- **Phone:** {phone}
- **Email:** {email}
- **LinkedIn:** {linkedin url or handle}
- **Availability:** Fully remote, {region} time zones
- **Languages:** {languages}

---

## Education (fixed)
**{Degree}**
{Institution}  ·  {City, Country}  ·  {Start} – {End}

---

## Fixed Skill Baselines
These appear in every resume. List ALL technical skills found, grouped by category.
Use this format:
- **Category:** skill1, skill2, skill3

---

## Work History
For EACH job, use this format:

### {N}. {Company Name}
**Location:** {city, state, country (remote/hybrid/onsite)}
**Type:** {Contract Full-Time / Full-Time / Part-Time / etc.}
**Period:** {Month Year} – {Month Year or Present}
**Default title:** {Job title}

**What they did here (use to write bullets):**
- {detailed bullet about what was built, technology used, impact}
- {another bullet}
(Extract as many specific, detailed bullets as possible from the resume text)

---

## Key Capabilities (use when writing summaries and bullets)
- **{Category}:** {capabilities}

---

IMPORTANT RULES:
- Extract ONLY information present in the resume — never invent or assume
- Calculate total years of experience from earliest job start to present
- Be as specific and detailed as possible in the work history bullets
- Preserve all technology names exactly as written

RESUME TEXT:
{resume_text}"""


def call_claude(resume_text: str, api_key: str) -> str:
    client = anthropic.Anthropic(api_key=api_key)
    prompt = PROFILE_PROMPT.replace('{resume_text}', resume_text)

    print('  Calling Claude API to extract profile...')
    message = client.messages.create(
        model='claude-sonnet-4-6',
        max_tokens=4096,
        messages=[{'role': 'user', 'content': prompt}]
    )
    return message.content[0].text


# ─── FILE GENERATORS ─────────────────────────────────────────────────────────

def generate_prompt_md(profile_text: str) -> str:
    # Extract name and contact from profile for personalization
    name_m    = re.search(r'\*\*Full name:\*\*\s*(.+)', profile_text)
    email_m   = re.search(r'\*\*Email:\*\*\s*(.+)',     profile_text)
    phone_m   = re.search(r'\*\*Phone:\*\*\s*(.+)',     profile_text)
    li_m      = re.search(r'\*\*LinkedIn:\*\*\s*(.+)',  profile_text)

    name  = name_m.group(1).strip()  if name_m  else 'Developer'
    email = email_m.group(1).strip() if email_m else ''
    phone = phone_m.group(1).strip() if phone_m else ''
    li    = li_m.group(1).strip()    if li_m    else ''

    return f"""# Resume Generation Prompt
# When given a JD, follow this workflow exactly. No explanation needed unless asked.

==================================================
REFERENCE FILES (read all before writing anything)
==================================================

1. _profile.md          — Developer's full background, work history, and capabilities
2. resume_prompt.md     — Mandatory rules, cloud stack definitions, timeline rules, output requirements
3. _resume_format.md    — Exact .md structure the resume must follow

==================================================
WORKFLOW
==================================================

1. Read the JD — identify: role level, required stack, cloud platform, key requirements
2. Read the 3 reference files above
3. Read CV_BUILD_PATH from .env. Replace spaces with underscores in Company and Position names.
   Create folder: {{CV_BUILD_PATH}}\\{{Company_underscored}}\\{{Position_underscored}}\\
4. Write resume.md       — tailored to 95%+ JD match, following _resume_format.md exactly
5. Write cover_letter.md — 7 sentences, following the cover letter format in _resume_format.md

PDF CONVERSION (use python directly — do NOT use convert.bat):
  python convert.py "{{CV_BUILD_PATH}}\\{{Company_underscored}}\\{{Position_underscored}}\\resume.md"
  python convert.py "{{CV_BUILD_PATH}}\\{{Company_underscored}}\\{{Position_underscored}}\\cover_letter.md"

OUTPUT FILES (name extracted automatically from _profile.md):
  {{name}}_CV.pdf           — A3, Style A
  {{name}}_CoverLetter.pdf  — B5, single page

==================================================
RESUME RULES
==================================================

- NEVER change the years of experience — always use the exact value from _profile.md
- Every bullet: WHAT was built → HOW it was built → BUSINESS or SYSTEM IMPACT
- Only use technologies and experience from _profile.md — never invent
- No symbols in bullet text: no —, <, >, +, →, ·
- Do not overuse metric values
- Every senior role must include: testing strategy, CI/CD, cloud services
- Strictly follow AI/LLM timeline rules (LangChain 2023+, LangGraph 2024+, LangFuse 2024+)

==================================================
COVER LETTER FORMAT (always exactly 7 sentences)
==================================================

Dear Hiring Manager,

[Sentence 1 — who the developer is + years of experience + top matching skill]
[Sentence 2 — most relevant technical match to this JD]
[Sentence 3 — second technical strength or achievement]
[Sentence 4 — third strength or specific JD requirement met]
[Sentence 5 — soft skill: leadership / remote-ready / English / cross-functional]
[Sentence 6 — why this company / role specifically]
[Sentence 7 — call to action]

Best regards,
{name}
Phone: {phone}
Email: {email}
LinkedIn: {li}

==================================================
JOB BID Q&A RULE
==================================================

If asked a question during a job application process, answer in 2-3 sentences max.
Clear, direct, no fluff.
"""


# ─── MAIN ────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description='CV Builder setup — generate config folder from existing resume'
    )
    parser.add_argument('resume', help='Path to existing resume (PDF or DOCX)')
    parser.add_argument('--output', '-o', required=True,
                        help='Output config folder path (will be created if needed)')
    args = parser.parse_args()

    resume_path = Path(args.resume).resolve()
    output_dir  = Path(args.output).resolve()
    script_dir  = Path(__file__).parent.resolve()

    if not resume_path.exists():
        print(f'ERROR: File not found: {resume_path}')
        sys.exit(1)

    api_key = os.environ.get('ANTHROPIC_API_KEY', '').strip()
    if not api_key:
        print('ERROR: ANTHROPIC_API_KEY environment variable is not set.')
        print('Set it with:  $env:ANTHROPIC_API_KEY = "sk-ant-..."')
        sys.exit(1)

    print(f'\nCV Builder Setup')
    print(f'  Resume : {resume_path.name}')
    print(f'  Output : {output_dir}')
    print()

    # Step 1: Extract text
    resume_text = extract_text(resume_path)
    if not resume_text.strip():
        print('ERROR: Could not extract any text from the resume file.')
        sys.exit(1)
    print(f'  Extracted {len(resume_text)} characters from resume.')

    # Step 2: Generate _profile.md via Claude
    profile_md = call_claude(resume_text, api_key)
    print(f'  Profile generated ({len(profile_md)} characters).')

    # Step 3: Create output folder
    output_dir.mkdir(parents=True, exist_ok=True)

    # Step 4: Write _profile.md
    (output_dir / '_profile.md').write_text(profile_md, encoding='utf-8')
    print('  Written: _profile.md')

    # Step 5: Copy shared template files from script directory
    # CSS files: use personal version if present, otherwise fall back to *.starter.css
    for fname in ('_resume_format.md', 'resume_prompt.md',
                  '_style_a.css', '_style_coverletter.css'):
        src = script_dir / fname
        if not src.exists() and fname.endswith('.css'):
            src = script_dir / fname.replace('.css', '.starter.css')
        dst = output_dir / fname
        if src.exists():
            label = f'{src.name} → {fname}' if src.name != fname else fname
            shutil.copy2(src, dst)
            print(f'  Copied : {label}')
        else:
            print(f'  WARNING: {fname} not found in {script_dir} — skipped.')

    # Step 6: Generate personalized GenerateResumePrompt.md
    prompt_md = generate_prompt_md(profile_md)
    (output_dir / 'GenerateResumePrompt.md').write_text(prompt_md, encoding='utf-8')
    print('  Written: GenerateResumePrompt.md')

    print(f'\nSetup complete. Config folder ready at:')
    print(f'  {output_dir}')
    print()
    print('Next steps:')
    print('  1. Review and adjust _profile.md if needed')
    print('  2. Customize _style_a.css to match your preferred style')
    print('  3. Run: convert.bat <resume.md> --config "' + str(output_dir) + '"')


if __name__ == '__main__':
    main()
